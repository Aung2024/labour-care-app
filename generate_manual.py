import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level):
    heading = doc.add_heading(text, level)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    return p

def add_alert(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"⚠️ သတိပြုရန် (Alert): {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 0, 0) # Red color
    return p

def add_placeholder(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ဤနေရာတွင် ဓာတ်ပုံထည့်ရန် - Insert Screenshot Here]")
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128) # Gray color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.5 * level)
    run = p.add_run(text)
    return p

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('m-MNCH Care အသုံးပြုသူ လမ်းညွှန်စာအုပ် (User Manual)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('ဤလမ်းညွှန်စာအုပ်သည် သားဖွားဆရာမများအတွက် m-MNCH Care အက်ပ်ကို အလွယ်တကူ အသုံးပြုနိုင်ရန် အသေးစိတ် ရှင်းလင်းထားသော လမ်းညွှန်စာအုပ်ဖြစ်ပါသည်။').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_placeholder(doc)
    doc.add_page_break()

    # 1. Login
    add_heading(doc, '၁။ အက်ပ်အတွင်းသို့ ဝင်ရောက်ခြင်း (Login)', 1)
    doc.add_paragraph('အက်ပ်ကို စတင်အသုံးပြုရန် မိမိ၏ အကောင့်ထဲသို့ ဝင်ရောက်ရပါမည်။')
    add_placeholder(doc)
    add_bullet(doc, 'Email (အီးမေးလ်) နှင့် Password (စကားဝှက်) အကွက်များတွင် မိမိအတွက် သတ်မှတ်ပေးထားသော အီးမေးလ်နှင့် စကားဝှက်ကို မှန်ကန်စွာ ရိုက်ထည့်ပါ။')
    add_bullet(doc, 'နောက်တစ်ကြိမ် အက်ပ်ဖွင့်တိုင်း ပြန်လည်ဝင်ရောက်စရာမလိုဘဲ အသင့်ဝင်ပြီးသားဖြစ်နေစေရန် "Keep me logged in" အကွက်လေးကို အမှန်ခြစ် (Tick) ပေးထားနိုင်ပါသည်။')
    add_bullet(doc, 'ထို့နောက် Login ခလုတ်ကို နှိပ်လိုက်ပါက Home (ပင်မစာမျက်နှာ) သို့ ရောက်ရှိသွားမည်ဖြစ်ပါသည်။')
    doc.add_page_break()

    # 2. Patient Registration
    add_heading(doc, '၂။ လူနာအသစ် မှတ်ပုံတင်ခြင်း (Patient Registration)', 1)
    doc.add_paragraph('လူနာအသစ်တစ်ဦးကို မှတ်ပုံတင်ရန်အတွက် Home (ပင်မစာမျက်နှာ) မှ "Register" ခလုတ်ကို နှိပ်ပါ။ (မှတ်ချက် - လူနာအသစ် မမှတ်ပုံတင်မီ လူနာစာရင်းတွင် ရှိမရှိ အရင်စစ်ဆေးပါ)။')
    add_placeholder(doc)
    add_bullet(doc, 'အနီရောင်ကြယ်ပွင့် (*) ပြထားသော အကွက်များအားလုံးကို မဖြစ်မနေ ဖြည့်စွက်ရပါမည်။')
    add_bullet(doc, 'Patient Name (လူနာအမည်): လူနာ၏အမည်ကို မှန်ကန်စွာရိုက်ထည့်ပါ။')
    add_bullet(doc, 'Age (အသက်): လူနာ၏အသက်ကို ဂဏန်းဖြင့် ထည့်သွင်းပါ။')
    add_alert(doc, 'အကယ်၍ အသက်သည် ၁၈ နှစ်အောက် သို့မဟုတ် ၃၅ နှစ်အထက် ဖြစ်ပါက Alert (သတိပေးချက်) ပြမည်ဖြစ်ပြီး၊ ထိုလူနာကို ပိုမိုကောင်းမွန်သော ဆေးရုံ/ဆေးခန်းသို့ လွှဲပြောင်းမွေးဖွားရန် အကြံပြုသင့်ပါသည်။')
    add_bullet(doc, 'Gravida (ကိုယ်ဝန်ဆောင်အကြိမ်) နှင့် Parity (မွေးဖွားအကြိမ်) ကို ဖြည့်စွက်ပါ။')
    add_alert(doc, 'အကယ်၍ ကလေး ၁ ယောက် သို့မဟုတ် ၁ ယောက်ထက်ပိုရှိပါက (Parity > 0)၊ နောက်ဆုံးကလေး၏ အသက် (Age of last child) ကို ဖြည့်ရပါမည်။ ကလေးအသက်သည် ၁ နှစ်အတွင်းဖြစ်ပါက Alert ပြမည်ဖြစ်ပြီး အထူးဂရုစိုက်ရန် လိုအပ်ပါသည်။')
    add_bullet(doc, 'Phone number (ဖုန်းနံပါတ်)၊ Address (နေရပ်လိပ်စာ) အစရှိသည်တို့ကို အပြည့်အစုံ ဖြည့်စွက်ရန် အကြံပြုပါသည်။')
    doc.add_paragraph('အချက်အလက်များ ဖြည့်စွက်ပြီးပါက "Save" (သိမ်းမည်) ခလုတ်ကို နှိပ်ပါ။ ထို့နောက် Patient Consent (လူနာသဘောတူညီချက်) စာမျက်နှာသို့ ရောက်ရှိပါမည်။')
    add_placeholder(doc)
    add_bullet(doc, 'ဤနေရာတွင် လူနာ၏ ကိုယ်ရေးကိုယ်တာနှင့် ကျန်းမာရေး အချက်အလက်များကို အက်ပ်တွင် သိမ်းဆည်းရန်အတွက် သဘောတူညီချက်ရယူရပါမည်။ လူနာအား ရှင်းပြပြီးပါက "Verbally agreed" (နှုတ်ဖြင့် သဘောတူညီချက်ရယူပြီး) ကို အမှန်ခြစ်၍သော်လည်းကောင်း၊ လူနာကိုယ်တိုင် လက်မှတ်ရေးထိုးခိုင်း၍သော်လည်းကောင်း လုပ်ဆောင်နိုင်ပါသည်။')
    doc.add_page_break()

    # 3. Select Patient
    add_heading(doc, '၃။ လူနာရှာဖွေခြင်းနှင့် ရွေးချယ်ခြင်း (Select Patient)', 1)
    doc.add_paragraph('Home စာမျက်နှာမှ "Select Patient" (လူနာရွေးချယ်မည်) ကို နှိပ်ပါက လူနာစာရင်း စာမျက်နှာသို့ ရောက်ရှိပါမည်။')
    add_placeholder(doc)
    add_bullet(doc, 'Search (ရှာဖွေရန်) အကွက်တွင် လူနာ၏ အမည် သို့မဟုတ် ဖုန်းနံပါတ်ကို ရိုက်ထည့်၍ အလွယ်တကူ ရှာဖွေနိုင်ပါသည်။')
    add_bullet(doc, 'Filter (စစ်ထုတ်ရန်) ခလုတ်ကို အသုံးပြု၍ လူနာများ၏ အခြေအနေအလိုက် (ဥပမာ- ကိုယ်ဝန်ဆောင်၊ မွေးဖွားပြီး) ခွဲခြားကြည့်ရှုနိုင်ပါသည်။')
    add_bullet(doc, 'Status (အခြေအနေ) များကို လူနာတစ်ဦးချင်းစီ၏ ကတ် (Card) ပေါ်တွင် ပြသထားမည်ဖြစ်ပြီး၊ ထိုလူနာသည် ကိုယ်ဝန်ဆောင်အဆင့်လား၊ မွေးဖွားပြီးအဆင့်လား ဆိုသည်ကို ချက်ချင်း သိရှိနိုင်ပါသည်။')
    add_alert(doc, 'လူနာအမည်ဘေးတွင် အနီရောင် Alert icon (သတိပေး သင်္ကေတ) လေးများ ပြနေပါက ထိုလူနာသည် အထူးဂရုစိုက်ရန် လိုအပ်သော "High Risk" လူနာ ဖြစ်သည်ကို သတိပြုရန်ဖြစ်သည်။ ၎င်းကို နှိပ်၍ မည်သည့်အန္တရာယ်ရှိသည်ကို အသေးစိတ် ကြည့်ရှုနိုင်ပါသည်။')
    add_bullet(doc, 'လူနာအချက်အလက်များကို ပြင်ဆင်လိုပါက "Edit" (ပြင်ဆင်မည်) ခလုတ်ကို နှိပ်၍ ပြန်လည်ပြင်ဆင်နိုင်ပါသည်။ လူနာအမည်ကို နှိပ်လိုက်ပါက ထိုလူနာ၏ Patient Care Hub (လူနာစောင့်ရှောက်မှုဗဟိုချက်) သို့ ရောက်ရှိမည်ဖြစ်ပါသည်။')
    doc.add_page_break()

    # 4. ANC Visit
    add_heading(doc, '၄။ ကိုယ်ဝန်ဆောင်စောင့်ရှောက်မှု (ANC Visit Form)', 1)
    doc.add_paragraph('Patient Care Hub မှ "ANC" ကို ရွေးချယ်ပြီး "Record New ANC Visit" (မှတ်တမ်းအသစ် ဖြည့်သွင်းရန်) ကို နှိပ်ပါ။ အောက်ပါအချက်အလက်များကို အသေးစိတ် ဖြည့်သွင်းရပါမည်။')
    add_placeholder(doc)
    add_bullet(doc, 'Visit Number (ပြသသည့်အကြိမ်): အက်ပ်မှ အလိုအလျောက် တွက်ချက်ပေးမည်ဖြစ်သည်။')
    add_bullet(doc, 'Date of Visit (လာရောက်ပြသသည့် ရက်စွဲ): ယနေ့ရက်စွဲကို ရွေးချယ်ပါ။')
    add_bullet(doc, 'Other Visits (အခြား စောင့်ရှောက်မှု): အခြားဆေးခန်း/ပုဂ္ဂလိကတွင် ပြသဖူးပါက ဖြည့်စွက်ပါ။')
    add_bullet(doc, 'LMP Status & EDD (နောက်ဆုံးရာသီပေါ်ရက် နှင့် ခန့်မှန်းမွေးဖွားရက်): LMP ရက်စွဲကို ထည့်လိုက်ပါက EDD ကို အက်ပ်မှ အလိုအလျောက် တွက်ပေးပါမည်။ LMP မမှတ်မိပါက Ultrasound အဖြေအရ ကိုယ်ဝန်ဆောင်ကာလ (Manual Gestational Age) ကို ထည့်သွင်းပါ။')
    add_bullet(doc, 'Weight (ကိုယ်အလေးချိန်) နှင့် Height (အရပ်): ကီလိုဂရမ် (Kg) နှင့် ထည့်သွင်းပါ။')
    add_bullet(doc, 'Pulse Rate (သွေးခုန်နှုန်း)၊ Blood Pressure (သွေးပေါင်ချိန်) နှင့် Temperature (ကိုယ်အပူချိန်) တိုင်းတာချက်များကို ဂဏန်းဖြင့် ထည့်သွင်းပါ။')
    add_alert(doc, 'သွေးပေါင်ချိန် (Blood Pressure) သည် ပုံမှန်ထက် မြင့်နေပါက အက်ပ်မှ ချက်ချင်း အနီရောင်ဖြင့် သတိပေးမည်ဖြစ်ပါသည်။')
    add_bullet(doc, 'Anemia (သွေးအားနည်းခြင်း)၊ Edema (ခြေလက်ဖောယောင်ခြင်း): အောက်မျက်ခမ်းနှင့် ခြေလက်များကို ကြည့်ရှုစစ်ဆေးပြီး "Yes" (ရှိသည်) သို့မဟုတ် "No" (မရှိပါ) ကို ရွေးချယ်ပါ။')
    add_bullet(doc, 'Co-infection symptoms (ကူးစက်ရောဂါ လက္ခဏာများ - TB/Malaria/HIV): ရှိပါက ထည့်သွင်းပါ။')
    add_bullet(doc, 'GBV (အကြမ်းဖက်ခံရမှု လက္ခဏာများ):')
    add_bullet(doc, 'ခန္ဓာကိုယ်တွင် အနာတရများရှိခြင်း (Bruise, cuts)', 1)
    add_bullet(doc, 'ပင်ပန်းနွမ်းနယ်ခြင်း၊ ဖြူဖျော့ခြင်း (Fatigue, pallor)', 1)
    add_bullet(doc, 'စိတ်ဖိစီးမှုနှင့် ကြောက်ရွံ့မှုများ (Worry, stress, fear)', 1)
    add_bullet(doc, 'Fundal Height (သားအိမ်အမြင့် - cm) နှင့် Fetal Heart Rate (သန္ဓေသားနှလုံးခုန်နှုန်း): (ကိုယ်ဝန် ၂၀ ပတ်အထက်)')
    add_bullet(doc, 'Fetal Position & Presentation (သန္ဓေသားအနေအထားနှင့် ဦးတိုက်မှု): (ကိုယ်ဝန် ၃၆ ပတ်အထက်)')
    
    doc.add_paragraph('High Risk Factors (အန္တရာယ်ဖြစ်နိုင်ခြေရှိသော အချက်အလက်များ): အောက်ပါအချက်များထဲမှ ကိုက်ညီသည်များကို အမှန်ခြစ် (Tick) ပေးပါ။')
    add_placeholder(doc)
    add_bullet(doc, 'အသက် ၁၈ နှစ်အောက် သို့မဟုတ် ၄၀ နှစ်အထက် (Age <=18 or >=40)', 1)
    add_bullet(doc, 'ဆီးချိုရောဂါ (Diabetes)', 1)
    add_bullet(doc, 'သွေးပေါင်ချိန်တက်ခြင်း (Hypertension)', 1)
    add_bullet(doc, 'နှလုံး/ကျောက်ကပ်ရောဂါ (Heart/Kidney Disease)', 1)
    add_bullet(doc, 'ကိုယ်ဝန်ဆိပ်တက်ခြင်း (Pre-eclampsia)', 1)
    add_bullet(doc, 'ယခင်ကိုယ်ဝန်တွင် နောက်ဆက်တွဲပြဿနာရှိဖူးခြင်း (Previous complications)', 1)
    add_bullet(doc, 'အမွှာကိုယ်ဝန် (Twins pregnancy)', 1)
    add_bullet(doc, 'အချင်းရှေ့ရောက်ခြင်း (Placenta previa)', 1)
    add_alert(doc, 'အထက်ပါ High Risk အချက်တစ်ချက်ခုကို ရွေးချယ်လိုက်ပါက လူနာသည် "High Risk Patient" အဖြစ် သတ်မှတ်ခံရပြီး အထူးစောင့်ကြည့်လူနာစာရင်းသို့ ရောက်ရှိသွားမည်ဖြစ်သည်။')

    doc.add_paragraph('Danger Signs (အန္တရာယ်လက္ခဏာများ): ကိုက်ညီပါက ရွေးချယ်ပါ။')
    add_placeholder(doc)
    add_bullet(doc, 'မိန်းမကိုယ်မှ သွေးဆင်းခြင်း (Vaginal Bleeding)', 1)
    add_bullet(doc, 'တက်ခြင်း/သတိလစ်ခြင်း (Convulsions)', 1)
    add_bullet(doc, 'ခေါင်းကိုက်ခြင်း/အမြင်မှုန်ခြင်း/အန်ခြင်း (Severe Headache, Blurred vision)', 1)
    add_bullet(doc, 'ဗိုက်အလွန်နာခြင်း (Severe Abdominal Pain)', 1)
    add_bullet(doc, 'ဖျားခြင်း (Fever) နှင့် အသက်ရှူမြန်ခြင်း (Fast breathing)', 1)
    add_bullet(doc, 'ခန့်မှန်းမွေးဖွားရက်ထက် ၇ ရက်ကျော်ခြင်း (Post Date)', 1)
    add_alert(doc, 'အန္တရာယ်လက္ခဏာများ ပြသနေပါက လိုအပ်သလို အရေးပေါ်ကုသမှုပေးရန်နှင့် လွှဲပြောင်းရန် (Transfer) စဉ်းစားရပါမည်။')

    doc.add_paragraph('Medications (ဆေးဝါးကုသမှုများ):')
    add_bullet(doc, 'သံဓာတ်နှင့် ဖောလစ်အက်ဆစ် (Iron Folic Acid)၊ ဗီတာမင်ဘီဝမ်း၊ မေးခိုင်ကာကွယ်ဆေး (Tetanus) အစရှိသည်တို့အတွက် ပေးသည့်အကြိမ်ရေ (Frequency) နှင့် ဆေးပြားအရေအတွက် (Total tablets) ကို ဖြည့်စွက်ပါ။')
    add_bullet(doc, 'Next Visit Date (နောက်တစ်ကြိမ် ပြန်ပြရက်) ကို အက်ပ်မှ အလိုအလျောက် တွက်ချက်ပေးမည်ဖြစ်ပြီး၊ မိမိစိတ်ကြိုက် ပြင်ဆင်လိုပါက "Adjust next visit date manually" ကို ရွေးချယ်၍ ပြင်ဆင်နိုင်ပါသည်။')
    add_bullet(doc, 'အဆုံးတွင် ဆရာမ၏ အမည်အတိုကောက် (Initial) ကို ထည့်သွင်းပြီး Save ခလုတ်ကို နှိပ်ပါ။')
    doc.add_page_break()

    # 5. Lab Tests & ANC Report & Health Education
    add_heading(doc, '၅။ ဓာတ်ခွဲစမ်းသပ်မှုများ (Lab Tests) နှင့် အခြား', 1)
    doc.add_paragraph('Patient Care Hub မှတဆင့် အောက်ပါအချက်အလက်များကို ဆက်လက် မှတ်တမ်းတင်နိုင်ပါသည်။')
    add_placeholder(doc)
    add_bullet(doc, 'Lab Tests (ဓာတ်ခွဲစမ်းသပ်မှုများ): "Record New Test" ကို နှိပ်ပြီး ဆီးချို၊ သွေးအားနည်းရောဂါ၊ Syphilis၊ HIV နှင့် ဆီးစစ်ဆေးမှု ရလဒ်များကို ဖြည့်သွင်းပါ။ ပုံမှန်မဟုတ်သော ရလဒ်များ (ဥပမာ- Positive) ထွက်ပေါ်ပါက Alert ပြမည်ဖြစ်ပါသည်။')
    add_bullet(doc, 'ANC Report (ကိုယ်ဝန်ဆောင်အစီရင်ခံစာ): ယခင်လာပြခဲ့သော ANC အချက်အလက်များအားလုံးကို ဇယားများ၊ ဂရပ်များဖြင့် စုစည်းကြည့်ရှုနိုင်ပြီး မိခင်၏ ကျန်းမာရေးတိုးတက်မှုကို အကဲဖြတ်နိုင်ပါသည်။')
    add_bullet(doc, 'Health Education (ကျန်းမာရေးပညာပေးခြင်း): မိခင်အား အာဟာရ၊ အန္တရာယ်လက္ခဏာများ အကြောင်း ပညာပေးပြီးတိုင်း "Mark as completed" ကို နှိပ်၍ မှတ်သားထားပါ။')
    doc.add_page_break()

    # 6. LCG
    add_heading(doc, '၆။ သားဖွားခြင်းစောင့်ကြည့်ရေးဇယား (LCG)', 1)
    doc.add_paragraph('မွေးဖွားခြင်းအဆင့် စတင်ပါက LCG ကို အသုံးပြုရပါမည်။ Patient Care Hub မှ "Labour Care" ကို နှိပ်ပါ။')
    add_placeholder(doc)
    
    add_paragraph(doc, '၆.၁ LCG Setup (အစပျိုးခြင်း)', bold=True)
    add_bullet(doc, 'Labour Onset (မွေးဖွားခြင်းစတင်သည့်ရက်နှင့် အချိန်): သဘာဝအတိုင်း (Spontaneous) လား၊ ဆေးဖြင့်ဖြစ်စေခြင်း (Induced) လား ရွေးချယ်ပြီး အချိန်ကို ဖြည့်သွင်းပါ။')
    add_bullet(doc, 'Ruptured Membranes (ရေမွှာပေါက်ချိန်): အချိန်အတိအကျကို ထည့်သွင်းပါ။')

    add_paragraph(doc, '၆.၂ LCG Entry (အချက်အလက်များ ဖြည့်သွင်းခြင်း)', bold=True)
    add_placeholder(doc)
    add_bullet(doc, 'Active First Stage Start Time (သားဖွားခြင်းတက်ကြွအဆင့် စတင်ချိန်): ဤအချိန်ကို ထည့်သွင်းလိုက်ပါက အက်ပ်မှ သင့်အတွက် အချိန်ကော်လံ (Time columns) များကို အလိုအလျောက် တွက်ချက်ဖန်တီးပေးသွားမည် ဖြစ်ပါသည်။ ဤအချိန်သည် သားအိမ်ခေါင်း ၅ စင်တီမီတာ ကျယ်သော အချိန်ဖြစ်ရပါမည်။')
    add_bullet(doc, 'Cervix Plot (သားအိမ်ခေါင်းအကျယ်) နှင့် Descent Plot (ကလေးဦးခေါင်းဆင်းသက်ခြင်း): ဇယားကွက်ပေါ်တွင် အမှတ်အသားလေးများ (Points) နှိပ်၍ ထည့်သွင်းသွားရပါမည်။')
    add_alert(doc, 'သားအိမ်ခေါင်းကျယ်ခြင်း မျဉ်းသည် Alert Line (သတိပေးမျဉ်း) သို့မဟုတ် Action Line (လုပ်ဆောင်ရန်မျဉ်း) ကို ကျော်လွန်သွားပါက ချက်ချင်း Alert ပြမည်ဖြစ်ပြီး အရေးပေါ် အရေးယူဆောင်ရွက်ရန် လိုအပ်ပါသည်။')
    
    add_paragraph(doc, 'အချိန်အလိုက် ဖြည့်သွင်းရမည့် အချက်များ -', bold=True)
    add_bullet(doc, 'Baseline FHR (သန္ဓေသားနှလုံးခုန်နှုန်း): ပုံမှန် 110-160 အတွင်း ရှိရပါမည်။')
    add_bullet(doc, 'Amniotic Fluid (ရေမွှာရည်): C (ကြည်လင်)၊ M (မစင်ပါ)၊ B (သွေးပါ)။ "M" သို့မဟုတ် "B" ဖြစ်ပါက သတိပေးချက် ပြပါမည်။')
    add_bullet(doc, 'Moulding (ဦးခေါင်းပုံသွင်းခံရခြင်း) နှင့် Caput (ဦးခေါင်းဖောင်းခြင်း): 0 (none), + (mild), ++ (moderate), +++ (severe) တို့မှ ရွေးချယ်ပါ။ "+++" ဖြစ်ပါက အန္တရာယ်ရှိပါသည်။')
    add_bullet(doc, 'Contractions per 10 min (၁၀ မိနစ်အတွင်း သားအိမ်ညှစ်ခြင်း) နှင့် Duration (ကြာချိန်)။')
    add_bullet(doc, 'Pulse (သွေးခုန်နှုန်း)၊ Systolic/Diastolic BP (သွေးပေါင်ချိန်)၊ Temperature (အပူချိန်) နှင့် Urine (ဆီး) အခြေအနေများကို အချိန်ကော်လံအလိုက် ဂရုတစိုက် ဖြည့်သွင်းပါ။')
    doc.add_page_break()

    # 7. LCG 3rd Stage
    add_heading(doc, '၇။ သားဖွားခြင်း တတိယအဆင့် (LCG 3rd Stage)', 1)
    doc.add_paragraph('ကလေးမွေးဖွားပြီးပါက Patient Care Hub သို့ပြန်သွားပြီး "LCG 3rd Stage" ကို နှိပ်ပါ။ ဤသည်မှာ အချင်းကျသည့် အဆင့်ကို တစ်ကြိမ်သာ မှတ်တမ်းတင်ရသည့် နေရာဖြစ်ပါသည်။')
    add_placeholder(doc)
    add_bullet(doc, 'အချင်းကျသည့်အချိန်၊ သွေးထွက်ရှိမှု ပမာဏ (Estimated Blood Loss) နှင့် အခြားပေးခဲ့သော ဆေးဝါးများကို မှတ်တမ်းတင်ပါ။')
    doc.add_page_break()

    # 8. Newborn Care
    add_heading(doc, '၈။ မွေးကင်းစကလေးစောင့်ရှောက်မှု (Newborn Care)', 1)
    doc.add_paragraph('Patient Care Hub မှ "Newborn Care" သို့ ဝင်ရောက်ပါ။ ကလေးမွေးပြီးချင်း ချက်ချင်း ပေးရမည့် စောင့်ရှောက်မှုများနှင့် နောက်ပိုင်းလာပြသော ပြသမှုများကို ဤနေရာတွင် ဖြည့်သွင်းရပါမည်။')
    
    add_paragraph(doc, '၈.၁ Immediate Newborn Care (မွေးမွေးချင်းစောင့်ရှောက်မှု)', bold=True)
    doc.add_paragraph('ဤစာမျက်နှာသည် တစ်ကြိမ်သာ ဖြည့်သွင်းရန် လိုအပ်ပါသည်။')
    add_placeholder(doc)
    add_bullet(doc, 'Apgar Score (၁ မိနစ် နှင့် ၅ မိနစ်) တွင် ရရှိသောအမှတ်များကို ထည့်သွင်းပါ။')
    add_bullet(doc, 'အောက်ပါ လုပ်ဆောင်ချက်များကို ပြီးမြောက်ပါက အမှန်ခြစ် (Tick) ပေးပါ -')
    add_bullet(doc, 'Thorough drying (သေချာစွာသုတ်သင်ခြင်း)', 1)
    add_bullet(doc, 'Spontaneous breathing (အလိုအလျောက် အသက်ရှူခြင်း) / Gasping or no breathing (အသက်ရှူခက်ခဲခြင်း/မရှူခြင်း)', 1)
    add_bullet(doc, 'Skin-to-skin contact (မိခင်နှင့် အရေပြားချင်းထိတွေ့စေခြင်း)', 1)
    add_bullet(doc, 'Delayed cord clamping (ချက်ကြိုးဖြတ်ခြင်းကို နောက်ဆုတ်ခြင်း)', 1)
    add_bullet(doc, 'Eye care (မျက်စိစောင့်ရှောက်မှု) နှင့် Cord care (ချက်ကြိုးစောင့်ရှောက်မှု)', 1)
    add_bullet(doc, 'Support early exclusive breastfeeding (မိခင်နို့စောစီးစွာတိုက်ကျွေးရန် ကူညီခြင်း)', 1)

    add_paragraph(doc, '၈.၂ Newborn Care Visit Form (ကလေးကျန်းမာရေးပြသမှု မှတ်တမ်း)', bold=True)
    doc.add_paragraph('ကလေး၏ နောက်ဆက်တွဲ စောင့်ရှောက်မှုများအတွက် "Record New Visit" ကို နှိပ်ပါ။')
    add_placeholder(doc)
    add_bullet(doc, 'Baby Name (ကလေးအမည်)၊ Gender (ကျား/မ)၊ Birth Time (မွေးဖွားချိန် - အလိုအလျောက် သော့ခတ်ထားမည်)၊ Weight (အလေးချိန်)၊ Length (အရှည်)၊ Head Circumference (ခေါင်းလုံးပတ်) တို့ကို ဖြည့်သွင်းပါ။')
    add_alert(doc, 'အကယ်၍ ကလေးသည် ပေါင်ချိန်မပြည့်ခြင်း (Weight < 2.5 kg) ဖြစ်ပါက KMC (Kangaroo Mother Care) လုပ်ဆောင်ရန် အက်ပ်မှ Alert ပြမည်ဖြစ်သည်။')
    add_bullet(doc, 'Eye Infection (မျက်စိပိုးဝင်ခြင်း)၊ Umbilical Cord Care (ချက်ကြိုးစောင့်ရှောက်မှု) နှင့် Anatomy abnormalities (ခန္ဓာကိုယ်ပုံပန်း ချို့ယွင်းချက်) ရှိ/မရှိ စစ်ဆေးပြီး ရွေးချယ်ပါ။')
    add_bullet(doc, 'Vaccines (ကာကွယ်ဆေးများ): Vitamin K, Hepatitis B, BCG ကာကွယ်ဆေးများ ထိုးပေးပြီးပါက အမှန်ခြစ် (Tick) ပေးပါ။')
    add_bullet(doc, 'Exclusive breastfeeding on demand (မိခင်နို့တစ်မျိုးတည်းကိုသာ လိုအပ်သလို တိုက်ကျွေးခြင်း) ကို အမှန်ခြစ်ပေးပါ။')
    
    add_paragraph(doc, 'Kangaroo Mother Care (KMC) အပိုင်း:', bold=True)
    add_bullet(doc, 'KMC ပြုလုပ်ရန် ရွေးချယ်ပါက "Yes" ကို နှိပ်ပြီး နေ့စဉ် KMC hour (KMC လုပ်သည့် အချိန်နာရီ)၊ Temperature (အပူချိန်)၊ Heart Rate (နှလုံးခုန်နှုန်း) နှင့် Respiration Rate (အသက်ရှူနှုန်း) ကို ဖြည့်စွက်ရပါမည်။ ပြုလုပ်ရန်မလိုပါက "No" ကို ရွေးချယ်ပြီး အကြောင်းရင်း (Reason category) ကို ရွေးပေးပါ။')
    
    add_paragraph(doc, 'Danger Signs (အန္တရာယ်လက္ခဏာများ): ကိုက်ညီပါက အမှန်ခြစ်ပါ။', bold=True)
    add_bullet(doc, 'Umbilical infection (ချက်တိုင် သွေးယို/ပြည်တည်ခြင်း/နီရဲခြင်း)', 1)
    add_bullet(doc, 'Skin infection (အရေပြား ပြည်ဖုများပေါက်/ရောင်ရမ်းခြင်း)', 1)
    add_bullet(doc, 'Severe Jaundice (ခြေဖဝါး/လက်ဖဝါးများ ဝါနေခြင်း)', 1)
    add_bullet(doc, 'Fast breathing (အသက်ရှူမြန်ခြင်း/ ရင်ဘတ်အောင့်ချိုင့်ဝင်ခြင်း)', 1)
    add_bullet(doc, 'Convulsion / Fit (တက်ခြင်း)', 1)
    add_bullet(doc, 'Fever / Hypothermia (ဖျားခြင်း/ ကိုယ်အေးစက်နေခြင်း)', 1)
    add_bullet(doc, 'Not feeding well (နို့ လုံးဝ မစို့လျှင်)', 1)
    add_bullet(doc, 'Very Weak / Unconscious (ခြေလက် လှုပ်ရှားမှု မရှိလျှင်/ သတိလစ်လျှင်)', 1)
    add_alert(doc, 'အထက်ပါ လက္ခဏာများတွေ့ရှိပါက အရေးပေါ်ကုသမှုပေးရန်နှင့် ဆေးရုံသို့ လွှဲပြောင်းရန် (Transfer) လိုအပ်ပါသည်။')
    doc.add_page_break()

    # 9. PNC
    add_heading(doc, '၉။ မီးဖွားပြီးစောင့်ရှောက်မှု (Postpartum/PNC Form)', 1)
    doc.add_paragraph('Patient Care Hub မှ "Postpartum Care" သို့ ဝင်ရောက်ပြီး "Record New PNC Visit" ကို နှိပ်ပါ။')
    add_placeholder(doc)
    add_bullet(doc, 'Visit Date (ပြသသည့်ရက်စွဲ) နှင့် Delivered Date & Time (မွေးဖွားသည့်ရက်စွဲနှင့်အချိန်) ကို စစ်ဆေးပါ။ မွေးဖွားချိန်ကို Newborn Form မှ အလိုအလျောက် ယူထားမည်ဖြစ်သည်။')
    add_bullet(doc, 'Pulse Rate, Blood Pressure, Temperature (သွေးခုန်နှုန်း၊ သွေးပေါင်ချိန်၊ ကိုယ်အပူချိန်) ကို တိုင်းတာဖြည့်သွင်းပါ။')
    add_bullet(doc, 'Breast Engorgement (ရင်သားတင်းမာမှု)၊ Breast Milk Production (နို့ထွက်ခြင်း) နှင့် Nipple Condition (နို့သီးခေါင်းအခြေအနေ) များကို သေချာစွာ စစ်ဆေးဖြည့်သွင်းပါ။')
    add_bullet(doc, 'Uterine Involution (သားအိမ်လုံးမာခြင်း/ဝပ်ခြင်း): ပုံမှန်ဟုတ်/မဟုတ် ရွေးချယ်ပါ။')
    add_bullet(doc, 'Vaginal Discharge / Lochia (မိန်းမကိုယ်မှ သွေးဆင်းခြင်း) နှင့် Discharge Smell (အနံ့) များကို စစ်ဆေးပါ။ အနံ့ဆိုး (Foul smell) ထွက်ပါက ပိုးဝင်နေနိုင်သဖြင့် သတိပြုပါ။')
    add_bullet(doc, 'Episiotomy Wound Healing (မွေးလမ်းကြောင်းအနာ/ခွဲစိတ်အနာ အခြေအနေ) ကို စစ်ဆေးပါ။')
    add_bullet(doc, 'Urination (ဆီးသွားခြင်း) နှင့် Constipation (ဝမ်းချုပ်ခြင်း) ပြဿနာရှိ/မရှိ မေးမြန်းပါ။')
    add_bullet(doc, 'Maternal Mental Health Assessment (စိတ်ကျန်းမာရေးဆိုင်ရာ ဆန်းစစ်ခြင်း): မိခင်၏ စိတ်ပိုင်းဆိုင်ရာ အခြေအနေများကို မေးမြန်းမှတ်တမ်းတင်ပါ။')
    add_bullet(doc, 'Vitamins (အားဆေးများ): Vitamin B Complex, Vitamin A, Iron Folic Acid များ တိုက်ကျွေးပါက အမှန်ခြစ်ပေးပါ။')
    add_bullet(doc, 'Contraception (သားဆက်ခြားနည်းလမ်းများ): မွေးပြီး ၄၂ ရက် နောက်ပိုင်း သားဆက်ခြားနည်းလမ်းများ (Injection/Pills/Condom/IUD/Implant) ဆွေးနွေးပေးခဲ့ပါက အမှန်ခြစ်ပေးပါ။')
    
    add_paragraph(doc, 'PNC Danger Signs (မီးဖွားပြီး အန္တရာယ်လက္ခဏာများ):', bold=True)
    add_placeholder(doc)
    add_bullet(doc, 'Heavy vaginal bleeding (မွေးဖွားပြီးချိန် သွေးသွန်ခြင်း)', 1)
    add_bullet(doc, 'Severe headache, Blurred vision (ပြင်းထန်စွာ ခေါင်းကိုက်/အမြင်မှုန်ဝါးခြင်း)', 1)
    add_bullet(doc, 'Convulsions (တက်ခြင်း) / Fast breathing (အသက်ရှူခက်ခဲခြင်း)', 1)
    add_bullet(doc, 'Breast Engorgement Danger (ရင်သားတင်းမာခြင်း၊ နို့သီးခေါင်းယောင်/နီ/နာခြင်း)', 1)
    add_bullet(doc, 'Foul smelling vaginal discharge (အနံ့ဆိုးသော သွေး/အရည်ဆင်းခြင်း)', 1)
    add_alert(doc, 'အဆိုပါ လက္ခဏာများရှိပါက မိခင်၏ ကျန်းမာရေးအတွက် အလွန်အန္တရာယ်ရှိသဖြင့် ချက်ချင်း အရေးယူဆောင်ရွက်ပါ။')
    
    add_paragraph(doc, 'Outcome (ရလဒ်):', bold=True)
    add_bullet(doc, 'Maternal Outcome (မိခင်ရလဒ်): Alive (အသက်ရှင်) သို့မဟုတ် Dead (သေဆုံး) ကို ရွေးချယ်ပါ။ သေဆုံးပါက အကြောင်းရင်း (Obstetric သို့မဟုတ် Non-obstetric) ကို ရွေးချယ်ပါ။')
    add_bullet(doc, 'Newborn Outcome (ကလေးရလဒ်): Alive (အသက်ရှင်)၊ Stillbirth (အသေမွေး) သို့မဟုတ် Neonatal Death (မွေးကင်းစသေဆုံး) ကို ရွေးချယ်ပါ။')
    add_bullet(doc, 'Treatment Given (ကုသမှုပေးခြင်းများ) နှင့် Clinical Notes (အကြံပြုချက်များ) ကို ဖြည့်စွက်ပြီး "Save" နှိပ်ပါ။')
    doc.add_page_break()

    # 10. Dashboards & Trackers
    add_heading(doc, '၁၀။ Dashboard, Trackers နှင့် လူနာလွှဲပြောင်းခြင်း', 1)
    
    add_paragraph(doc, '၁၀.၁ Dashboard (ဒက်ရှ်ဘုတ်)', bold=True)
    doc.add_paragraph('အက်ပ်၏ ပင်မစာမျက်နှာ (Home) တွင်ရှိပြီး၊ တစ်ပတ်/တစ်လအတွင်း မှတ်ပုံတင်ထားသော လူနာအရေအတွက်၊ ANC/PNC ကြည့်ရှုမှု အရေအတွက်များကို အကျဉ်းချုပ် ပြသပေးပါသည်။ လစဉ်အစီရင်ခံစာများ (Monthly Reports) ကို ဤနေရာမှနေ၍ လွယ်ကူစွာ Print ထုတ်ယူနိုင်ပါသည်။')
    add_placeholder(doc)

    add_paragraph(doc, '၁၀.၂ High Risk Tracker (အန္တရာယ်ရှိလူနာများ စောင့်ကြည့်စစ်ဆေးခြင်း)', bold=True)
    doc.add_paragraph('ဘေးဘက် Menu မှ "High Risk Tracking" သို့ ဝင်ရောက်ပါ။ ဤစာမျက်နှာတွင် High Risk အဖြစ် သတ်မှတ်ခံထားရသော ကိုယ်ဝန်ဆောင်များ အားလုံးကို တစ်နေရာတည်းတွင် စုစည်းပြသထားမည်ဖြစ်သည်။ လူနာမည်မျှ အန္တရာယ်ရှိနေသည်ကို လွယ်ကူစွာ သိရှိနိုင်ပြီး ကြိုတင်ပြင်ဆင်မှုများ ပြုလုပ်ထားနိုင်ပါသည်။')
    add_placeholder(doc)

    add_paragraph(doc, '၁၀.၃ KMC Tracker (KMC စောင့်ကြည့်စစ်ဆေးခြင်း)', bold=True)
    doc.add_paragraph('ဘေးဘက် Menu မှ "KMC Tracking" သို့ ဝင်ရောက်ပါ။ ပေါင်မပြည့်သော ကလေးငယ်များအတွက် Kangaroo Mother Care (ရင်ခွင်ပိုက်စောင့်ရှောက်ခြင်း) ခံယူနေသော ကလေးစာရင်းကို ပြသပေးပါသည်။ ကလေးအလေးချိန် ပုံမှန်ရောက်ရှိသွားပါက KMC အောင်မြင်စွာ ပြီးမြောက်ကြောင်း မှတ်သားနိုင်ပါသည်။')
    add_placeholder(doc)

    add_paragraph(doc, '၁၀.၄ Transfer/Receive Patients (လူနာလွှဲပြောင်းခြင်းနှင့် လက်ခံခြင်း)', bold=True)
    doc.add_paragraph('ဘေးဘက် Menu မှ "Patient Transfers" သို့ ဝင်ရောက်ပါ။')
    add_bullet(doc, 'လူနာလွှဲပြောင်းရန် (Transfer out): Transfer Form (လွှဲပြောင်းလွှာ) တွင် လူနာ၏ လက်ရှိအခြေအနေ၊ ပေးခဲ့သော ကုသမှုများနှင့် လွှဲပြောင်းရသည့် အကြောင်းရင်းကို တိကျစွာ ဖြည့်စွက်၍ အခြားဆေးရုံ/ဆေးခန်းသို့ လွှဲပြောင်းပါ။')
    add_bullet(doc, 'လူနာလက်ခံရန် (Receive in): အခြားဆေးခန်းမှ မိမိထံသို့ လွှဲပြောင်းလာသော လူနာများကို ဤနေရာတွင် "Receive" ခလုတ်ကို နှိပ်၍ မိမိလူနာစာရင်းထဲသို့ ထည့်သွင်းလက်ခံနိုင်ပါသည်။')
    add_placeholder(doc)

    add_paragraph(doc, '၁၀.၅ Scoreboard (ရမှတ်နှင့် အဆင့်သတ်မှတ်ချက်)', bold=True)
    doc.add_paragraph('ဘေးဘက် Menu မှ "Leaderboard" သို့ ဝင်ရောက်ပါ။ သားဖွားဆရာမများ၏ လုပ်ငန်းဆောင်ရွက်မှုများကို ဂုဏ်ပြုသည့်အနေဖြင့် အမှတ်ပေးစနစ် ပါဝင်ပါသည်။ လူနာအသစ်မှတ်ပုံတင်ခြင်း၊ ANC/PNC ကြည့်ရှုပေးခြင်းတိုင်းအတွက် Points များ ရရှိမည်ဖြစ်ပြီး၊ မိမိ၏ အဆင့် (Rank) ကို အခြားလုပ်ဖော်ကိုင်ဖက်များနှင့် နှိုင်းယှဉ်ကြည့်ရှုနိုင်ပါသည်။')
    add_placeholder(doc)

    # Save Document
    doc.save('m-MNCH_Care_User_Manual.docx')
    print("Document successfully generated: m-MNCH_Care_User_Manual.docx")

if __name__ == '__main__':
    main()