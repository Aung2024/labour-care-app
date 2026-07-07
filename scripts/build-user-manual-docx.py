#!/usr/bin/env python3
"""Build Myanmar-only m-MNCH Care user manual Word document."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(Path(__file__).resolve().parent))
from manual_form_fields import apply_detailed_sections

OUT_PATH = ROOT / "docs" / "m-MNCH-Care-User-Manual.docx"
OUT_PATH_ALT = ROOT / "m-MNCH_Care_User_Manual.docx"
LOGO_PATH = ROOT / "icons" / "icon-512.png"
ASSETS_DIR = ROOT / "docs" / "user-manual-assets"
MOBILE_SHOT_WIDTH = Inches(2.35)
FONT_MM = "Myanmar Text"

# A5 book size (ISO 148 × 210 mm)
A5_WIDTH = Mm(148)
A5_HEIGHT = Mm(210)
# Binding: extra space on inner (left) edge for saddle/stitch or perfect binding
BIND_GUTTER = Inches(0.25)
BIND_INNER = Inches(0.85)
BIND_OUTER = Inches(0.55)
BIND_TOP = Inches(0.55)
BIND_BOTTOM = Inches(0.6)

# Book palette
COLOR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)
COLOR_PRIMARY_HEX = "1E40AF"
COLOR_ACCENT_HEX = "3B82F6"
COLOR_GOLD_HEX = "D97706"
COLOR_MUTED = RGBColor(0x47, 0x55, 0x69)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_run_font(run, size=11, bold=False, color=None, italic=False):
    run.font.name = FONT_MM
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_MM)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_vertical_align(cell, align="center"):
    tc_pr = cell._tc.get_or_add_tcPr()
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), align)
    tc_pr.append(valign)


def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def add_para(doc, text, size=11, bold=False, color=None, space_after=6, alignment=None, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_para_in_cell(cell, text, size=11, bold=False, color=None, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8):
    p = cell.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    for part, text in (
        ("begin", None),
        (None, "PAGE"),
        ("separate", None),
        (None, "1"),
        ("end", None),
    ):
        if part is None:
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = text
            run._r.append(instr)
        else:
            fld = OxmlElement("w:fldChar")
            fld.set(qn("w:fldCharType"), part)
            run._r.append(fld)
    set_run_font(run, size=9, color=COLOR_MUTED)


def clear_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    for paragraph in section.header.paragraphs:
        paragraph.clear()
    for paragraph in section.footer.paragraphs:
        paragraph.clear()


def set_page_size(section, *, landscape=False):
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = A5_HEIGHT
        section.page_height = A5_WIDTH
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = A5_WIDTH
        section.page_height = A5_HEIGHT


def configure_section(
    section,
    *,
    landscape=False,
    margins=None,
    gutter=None,
    header_footer=False,
):
    set_page_size(section, landscape=landscape)
    if margins is None:
        if landscape:
            margins = (BIND_TOP.inches, BIND_BOTTOM.inches, BIND_INNER.inches, BIND_OUTER.inches)
        else:
            margins = (0.5, 0.5, 0.65, 0.55)
    top, bottom, left, right = margins
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.gutter = gutter if gutter is not None else (BIND_GUTTER if landscape else Inches(0))
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    if header_footer:
        header = section.header.paragraphs[0]
        header.text = ""
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("m-MNCH Care အသုံးပြုသူ လမ်းညွှန်စာအုပ်")
        set_run_font(run, size=8, color=COLOR_MUTED)
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number_field(footer)


def start_page_numbering_at_one(section):
    sect_pr = section._sectPr
    pg_num = OxmlElement("w:pgNumType")
    pg_num.set(qn("w:start"), "1")
    sect_pr.append(pg_num)


def add_cover_page(doc):
    section = doc.sections[0]
    clear_header_footer(section)
    configure_section(section, landscape=False, gutter=Inches(0))

    content_width = A5_WIDTH - Inches(1.2)
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = content_width
    cell = table.rows[0].cells[0]
    shade_cell(cell, COLOR_PRIMARY_HEX)
    set_cell_vertical_align(cell, "center")
    set_cell_margins(cell, top=480, start=360, bottom=480, end=360)

    if LOGO_PATH.exists():
        logo_p = cell.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_after = Pt(12)
        run = logo_p.add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(0.95))

    add_para_in_cell(cell, "m-MNCH Care", size=26, bold=True, color=COLOR_WHITE, space_after=4)
    add_para_in_cell(cell, "အသုံးပြုသူ လမ်းညွှန်စာအုပ်", size=18, bold=True, color=COLOR_WHITE, space_after=12)

    rule = cell.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule_run = rule.add_run("━━━━━━━━━━━━━━━━━━━━━━")
    set_run_font(rule_run, size=10, color=RGBColor(0xF5, 0x9E, 0x0B))
    rule.paragraph_format.space_after = Pt(18)

    add_para_in_cell(
        cell,
        "သားဖွားဆရာမများနှင့် ကျန်းမာရေးဝန်ထမ်းများအတွက်\n"
        "အကောင့်ဖွင့်ခြင်းမှ လူနာမှတ်တမ်း၊ ANC၊ LCG၊ မွေးကင်းစ၊ PNC၊\n"
        "ကာကွယ်ဆေး၊ လွှဲပြောင်းမှု၊ အစီရင်ခံစာနှင့် Sync အထိ",
        size=10,
        color=RGBColor(0xDB, 0xEA, 0xFE),
        space_after=18,
    )

    add_para_in_cell(cell, "ဗားရှင်း ၂၀၂၆", size=10, bold=True, color=COLOR_WHITE, space_after=3)
    add_para_in_cell(cell, "A5 · Myanmar Edition", size=9, color=RGBColor(0xBF, 0xDB, 0xFE), space_after=24)

    add_para_in_cell(cell, "ကျန်းမာရေးဝန်ကြီးဌာန", size=11, bold=True, color=COLOR_WHITE, space_after=3)
    add_para_in_cell(cell, "Jhpiego Myanmar", size=11, color=RGBColor(0xDB, 0xEA, 0xFE), space_after=0)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_title_page(doc):
    for _ in range(2):
        doc.add_paragraph()

    add_para(
        doc,
        "m-MNCH Care",
        size=24,
        bold=True,
        color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=6,
    )
    add_para(
        doc,
        "အသုံးပြုသူ လမ်းညွှန်စာအုပ်",
        size=18,
        bold=True,
        color=COLOR_PRIMARY,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=14,
    )
    add_para(
        doc,
        "မိခင်နှင့်ကလေး ကျန်းမာရေး စောင့်ရှောက်မှု အက်ပ်အသုံးပြုရန် အဆင့်လိုက် လမ်းညွှန်",
        size=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=36,
    )
    add_para(doc, "ကျန်းမာရေးဝန်ကြီးဌာန", size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Jhpiego Myanmar", size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, color=COLOR_MUTED, space_after=24)
    add_para(doc, "ဗားရှင်း ၂၀၂၆", size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_front_matter_page(doc):
    add_heading(doc, "ဤစာအုပ်အကြောင်း", level=1)
    add_para(
        doc,
        "ဤလမ်းညွှန်စာအုပ်သည် m-MNCH Care app ကို နေ့စဉ်ဆေးခန်းလုပ်ငန်းတွင် အသုံးပြုရန် ရည်ရွယ်ထားသော လက်တွေ့ လမ်းညွှန်ဖြစ်ပါသည်။",
        size=11,
    )
    add_para(doc, "အသုံးပြုပုံ", size=12, bold=True, space_after=4)
    add_bullets(
        doc,
        [
            "ပုံကိုကြည့်ပါ — screenshot နေရာများတွင် app ပုံ ထည့်သွင်းထားပါသည်။",
            "အဆင့်ကိုဖတ်ပါ — နံပါတ်တပ်ထားသော အဆင့်များကို အစဉ်လိုက် လုပ်ဆောင်ပါ။",
            "app တွင် လုပ်ဆောင်ပါ — လမ်းညွှန်နှင့် app ကို ဘေးချင်းယှဉ်ကြည့်ပါ။",
        ],
    )
    add_para(doc, "ဤစာအုပ်ကို training၊ onboarding နှင့် နေ့စဉ်ကိုးကားချက်အဖြစ် အသုံးပြုနိုင်ပါသည်။", size=11)
    add_callout(
        doc,
        "မှတ်ချက်",
        "ကလင်နစ်အကြောင်းအရာများကို medical team က Word ထဲတွင် တိုက်ရိုက်ပြင်ဆင်နိုင်ပါသည်။",
        "info",
    )
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


_chapter_heading_count = 0


def add_heading(doc, text, level=1, page_break_before=False):
    global _chapter_heading_count
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=15, bold=True, color=COLOR_PRIMARY)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        if page_break_before:
            run.add_break(WD_BREAK.PAGE)
        _chapter_heading_count += 1
    elif level == 2:
        set_run_font(run, size=12, bold=True, color=RGBColor(0x11, 0x18, 0x27))
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, bold=True)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_steps(doc, items):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(4)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(3)


def add_callout(doc, label, text, kind="info"):
    colors = {
        "info": RGBColor(0x1E, 0x40, 0xAF),
        "warning": RGBColor(0xB4, 0x53, 0x09),
        "danger": RGBColor(0xB9, 0x1C, 0x1C),
        "success": RGBColor(0x05, 0x96, 0x69),
    }
    p = doc.add_paragraph()
    r1 = p.add_run(label + " ")
    set_run_font(r1, bold=True, color=colors.get(kind, colors["info"]))
    r2 = p.add_run(text)
    set_run_font(r2)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(8)


def add_screenshot_placeholder(doc, title, filename, caption):
    p = doc.add_paragraph()
    r = p.add_run(f"[ဤနေရာတွင် mobile screenshot ထည့်ရန်]\n{title}\nဖိုင်: {filename}")
    set_run_font(r, size=9, color=RGBColor(0x25, 0x63, 0xEB))
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    add_para(doc, caption, size=9, color=RGBColor(0x47, 0x55, 0x69))


def add_screenshot_cell_image(cell, title, filename, caption):
    img_path = ASSETS_DIR / filename
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if img_path.exists():
        run = p.add_run()
        run.add_picture(str(img_path), width=MOBILE_SHOT_WIDTH)
    else:
        run = p.add_run(f"[Mobile screenshot]\n{title}\n{filename}")
        set_run_font(run, size=8, color=RGBColor(0x25, 0x63, 0xEB))
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=8, color=RGBColor(0x47, 0x55, 0x69))


def add_screenshot_row(doc, shots, cols=2):
    """Place mobile screenshots side by side (A5 landscape). Up to 2 per row."""
    if not shots:
        return
    count = min(len(shots), cols)
    table = doc.add_table(rows=1, cols=count)
    table.autofit = False
    col_w = int((A5_HEIGHT - BIND_INNER - BIND_OUTER - BIND_GUTTER) / count)
    for ci in range(count):
        table.columns[ci].width = col_w
        title, fname, caption = shots[ci]
        add_screenshot_cell_image(table.cell(0, ci), title, fname, caption)
    if len(shots) > cols:
        add_screenshot_row(doc, shots[cols:], cols=cols)
    doc.add_paragraph()


def add_field_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True, size=8)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=8)
    doc.add_paragraph()


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)
    doc.add_paragraph()


MANUAL = [
    {
        "title": "စတင်အသုံးပြုခြင်း",
        "sections": [
            {
                "heading": "App အကြောင်း",
                "paragraphs": [
                    "m-MNCH Care app သည် မိခင်နှင့်ကလေး ကျန်းမာရေး စောင့်ရှောက်မှုအတွက် လူနာတစ်ဦး၏ ခရီးစဉ်ကို အစအဆုံး မှတ်တမ်းတင်နိုင်သော app ဖြစ်သည်။",
                ],
                "bullets": [
                    "လူနာအသစ် မှတ်ပုံတင်နိုင်သည်။",
                    "ANC, LCG, Newborn, PNC မှတ်တမ်းများ သိမ်းနိုင်သည်။",
                    "အစီရင်ခံစာများ ကြည့်ရှုပြီး ပရင့်ထုတ်နိုင်သည်။",
                    "Internet မရှိသေးသည့်အချိန်တွင်လည်း အချို့လုပ်ငန်းများကို Offline ဖြင့် ဆက်လုပ်နိုင်သည်။",
                ],
                "callouts": [("warning", "သတိပြုရန်", "လူနာ၏ အချက်အလက်များသည် လုံခြုံရေးအရေးကြီးပါသည်။ မိမိ password ကို မည်သူ့ကိုမျှ မပေးပါနှင့်။")],
                "screenshots": [
                    ("Login screen", "01-login.png", "App ဖွင့်သောအခါ Login စာမျက်နှာကို တွေ့ရပါမည်။"),
                    ("Home screen", "04-home.png", "Login ဝင်ပြီးနောက် Home screen မှ လုပ်ငန်းများကို စတင်ပါ။"),
                ],
            },
            {
                "heading": "လူနာခရီးစဉ် အကျဉ်းချုပ်",
                "paragraphs": [
                    "လူနာတစ်ဦးအတွက် အကြံပြုလုပ်ငန်းအစဉ်မှာ အောက်ပါအတိုင်း ဖြစ်ပါသည် — ရှိပြီးသားလူနာကို ရှာပါ၊ မရှိသေးပါက မှတ်ပုံတင်ပါ၊ Patient Care Hub မှ သက်ဆိုင်ရာ module ကိုရွေးပါ။",
                ],
                "steps": [
                    "Home → Select Patient for Care ဖြင့် လူနာရှာပါ (မှတ်ပုံတင်မီ အမြဲရှာပါ)။",
                    "မိခင်လူနာ — ANC → LCG → မွေးကင်းစ → PNC အစဉ်လိုက် မှတ်တမ်းတင်ပါ။",
                    "ကလေးလူနာ — Immunization နှင့် newborn follow-up မှတ်တမ်းများ သိမ်းပါ။",
                    "လိုအပ်ပါက Transfer ဖြင့် အတွင်းပိုင်း သို့မဟုတ် ပြင်ပ facility သို့ လွှဲပြောင်းပါ။",
                    "နေ့အဆုံးတွင် Sync လုပ်ပြီး မှတ်တမ်းများ cloud သို့ရောက်ကြောင်း စစ်ပါ။",
                ],
            },
        ],
    },
    {
        "title": "အကောင့်ဖွင့်ခြင်း",
        "sections": [
            {
                "heading": "အကောင့်အသစ် တောင်းဆိုခြင်း",
                "steps": [
                    "Login စာမျက်နှာတွင် Register သို့မဟုတ် Create Account ကိုနှိပ်ပါ။",
                    "အမည်၊ Email၊ ဖုန်းနံပါတ်၊ Role၊ Region၊ Township၊ Facility ကို ဖြည့်ပါ။",
                    "Password ကို မှန်ကန်စွာ ထည့်ပါ။ မှတ်မိလွယ်သော်လည်း လုံခြုံသော password ဖြစ်ရမည်။",
                    "Register ကိုနှိပ်ပါ။",
                    "အကောင့်ကို အုပ်ချုပ်သူက အတည်ပြုရပါမည်။ အတည်ပြုပြီးမှ Login ဝင်နိုင်ပါမည်။",
                ],
                "callouts": [("info", "မအောင်မြင်ပါက", "Email မှားခြင်း၊ Internet မရှိခြင်း၊ Facility မရွေးခြင်းတို့ကို စစ်ပါ။")],
                "screenshots": [
                    ("Account registration", "02-registration.png", "အကောင့်ဖွင့်ရာတွင် Role နှင့် Facility ကို မှန်မှန်ကန်ကန် ရွေးပါ။"),
                    ("Registration success", "03-registration-success.png", "အကောင့်တောင်းဆိုပြီးပါက အတည်ပြုမှုကို စောင့်ပါ။"),
                ],
            },
        ],
    },
    {
        "title": "Login နှင့် Consent",
        "sections": [
            {
                "heading": "Login ဝင်ခြင်း",
                "steps": [
                    "Email box တွင် မိမိ email ကို ရိုက်ပါ။",
                    "Password box တွင် password ကို ရိုက်ပါ။",
                    "Offline သုံးရန် လိုပါက Keep me logged in ကို အမှန်ခြစ်ပါ။",
                    "Login ကိုနှိပ်ပါ။",
                    "ပထမဆုံးအသုံးပြုပါက Provider Consent စာမျက်နှာပေါ်လာနိုင်ပါသည်။ စာကိုဖတ်ပြီး သဘောတူပါ။",
                ],
                "screenshots": [
                    ("Login screen", "01-login.png", "Login စာမျက်နှာတွင် Email, Password နှင့် Keep me logged in ကို စစ်ပါ။"),
                    ("Provider consent", "03-provider-consent.png", "Provider Consent ကို သဘောတူပြီးမှ app ကို ဆက်သုံးနိုင်ပါမည်။"),
                ],
            },
        ],
    },
    {
        "title": "Home Screen နှင့် Cards များ",
        "sections": [
            {
                "heading": "ပင်မစာမျက်နှာ အသုံးပြုပုံ",
                "steps": [
                    "Home တွင် မိမိအကောင့်၊ Role၊ Facility အချက်အလက်ကို စစ်ပါ။",
                    "မြန်မာ/အင်္ဂလိပ် ပြောင်းလိုပါက MM သို့မဟုတ် ENG ကိုနှိပ်ပါ။",
                    "လူနာသစ်အတွက် Patient Registration ကိုနှိပ်ပါ။",
                    "ရှိပြီးသားလူနာကို ရှာရန် Select Patient for Care ကိုနှိပ်ပါ။",
                    "Offline မှတ်တမ်းများရှိပါက Sync ကိုနှိပ်ပြီး Cloud သို့ပို့ပါ။",
                ],
            },
            {
                "heading": "Home Cards များ၏ အဓိပ္ပါယ်",
                "paragraphs": [
                    "Home screen တွင် မိမိ role အလိုက် card များကွဲပြားနိုင်ပါသည်။ အချို့ card များသည် Midwife အတွက်ဖြစ်ပြီး၊ အချို့မှာ TMO/Admin အတွက်သာ ပေါ်နိုင်ပါသည်။",
                ],
                "table": {
                    "headers": ["Card အမည်", "ဘာအတွက်သုံးသလဲ", "အသုံးပြုပုံ"],
                    "rows": [
                        ["Patient Registration\nလူနာမှတ်ပုံတင်ခြင်း", "လူနာအသစ်ကို စနစ်ထဲသို့ စတင်ထည့်ရန်", "Card ကိုနှိပ်ပါ၊ form ဖြည့်ပါ၊ consent ယူပြီး save လုပ်ပါ"],
                        ["Select Patient for Care\nလူနာရွေးချယ်ရန်", "ရှိပြီးသားလူနာကိုရှာပြီး care modules များသို့ဝင်ရန်", "Search/filter ဖြင့်လူနာရှာပါ၊ မှန်ကန်သော patient card ကိုနှိပ်ပါ"],
                        ["Patient Transfers\nလူနာ လွှဲပြောင်းမှုများ", "လူနာကို အခြား facility သို့လွှဲရန် သို့မဟုတ် ဝင်လာသော transfer request များလက်ခံရန်", "Open Transfers ကိုနှိပ်ပါ။ ဝင်လာသော request ရှိပါက badge ကိုစစ်ပါ"],
                        ["Analytics Dashboard\nဒက်ရှ်ဘုတ်", "Facility/လုပ်ငန်းဆောင်ရွက်မှု အချက်အလက်များကို charts, counts အဖြစ်ကြည့်ရန်", "View Dashboard ကိုနှိပ်ပါ။ Internet ရှိသောအချိန်တွင် အသုံးပြုပါ"],
                        ["Township Report\nမြို့နယ် အစီရင်ခံစာ", "TMO role အတွက် မြို့နယ်အဆင့် A4 printable report ထုတ်ရန်", "Generate Report ကိုနှိပ်ပြီး print/PDF ထုတ်ပါ"],
                        ["CME\nစဉ်ဆက်မပျက် လေ့လာသင်ယူမှု", "Mandatory/optional learning module များလေ့လာရန်", "Start Learning ကိုနှိပ်ပါ"],
                        ["High Risk Tracking\nအန္တရာယ်မြင့် စောင့်ကြည့်မှု", "Follow-up လိုသောအန္တရာယ်မြင့် လူနာများကြည့်ရန်", "Open High Risk Tracking ကိုနှိပ်ပါ"],
                        ["Scoreboard\nအမှတ်ပေးဇယား", "Midwife performance ranking ကြည့်ရန်", "View Leaderboard ကိုနှိပ်ပါ"],
                    ],
                },
                "callouts": [("warning", "သတိပြုရန်", "Dashboard, Township Report, Transfers, CME စသည့် card များသည် online-only ဖြစ်နိုင်ပါသည်။")],
                "screenshots": [
                    ("Home screen", "04-home.png", "နေ့စဉ်လုပ်ငန်းအများစုသည် Home မှ စတင်ပါသည်။"),
                    ("Home cards", "28-home-cards.png", "Role အလိုက် card များကို ဤနေရာတွင် screenshot ထည့်ပါ။"),
                ],
            },
        ],
    },
    {
        "title": "လူနာမှတ်ပုံတင်ခြင်း",
        "sections": [
            {
                "heading": "မိခင်လူနာ မှတ်ပုံတင်ခြင်း",
                "steps": [
                    "Home မှ Patient Registration ကိုနှိပ်ပါ။",
                    "လူနာအမည်၊ အသက် (၁၂ နှစ်နှင့် အထက်), ဖုန်းနံပါတ်၊ လိပ်စာကို ဖြည့်ပါ။",
                    "ကိုယ်ဝန်ဆိုင်ရာ အချက်အလက်များ — Gravida, Parity, LMP, EDD ကို ဖြည့်ပါ။",
                    "လိုအပ်သော field များကို မဖြည့်ပါက Save မလုပ်နိုင်ပါ။ အနီရောင် သတိပေးစာကို ဖတ်ပါ။",
                    "Save သို့မဟုတ် Continue ကိုနှိပ်ပါ။",
                    "Patient Consent စာမျက်နှာပေါ်လာပါက လူနာအားရှင်းပြပြီး သဘောတူချက်ယူပါ။",
                ],
                "callouts": [("success", "ကောင်းမွန်သော အလေ့အကျင့်", "လူနာမသွင်းမီ list တွင် ဖုန်းနံပါတ် သို့မဟုတ် အမည်ဖြင့် ရှာပါ။ Duplicate မဖြစ်စေရန် အရေးကြီးပါသည်။")],
            },
            {
                "heading": "ကလေးလူနာ မှတ်ပုံတင်ခြင်း",
                "steps": [
                    "အသက် ၁၂ နှစ်အောက် ထည့်ပါက ကလေးလူနာ အဖြစ် မှတ်ပုံတင်နိုင်ပါသည်။",
                    "မိခင်အမည်၊ မွေးဖွားရက် (date only)၊ ကျား/မ ကို ဖြည့်ပါ။",
                    "စနစ်ထဲတွင် မိခင်ရှိပါက လင့်ခ်ချိတ်ပါ။ မရှိပါက consent အဆင့်တွင် ဆက်လုပ်နိုင်ပါသည်။",
                    "ဤနေရာတွင် delivery notes မှ အလိုအလျောက် ဖန်တီးထားသော ကလေးကို ပြန်မှတ်ပုံတင်ရန် မလိုပါ။",
                ],
            },
            {
                "heading": "Duplicate စစ်ဆေးခြင်း",
                "paragraphs": [
                    "ဖုန်းနံပါတ် တူညီပါက မှတ်ပုံတင်ခြင်း မလုပ်နိုင်ပါ။ အမည်တူသော်လည်း warning ပြနိုင်ပါသည်။",
                ],
                "screenshots": [
                    ("Patient registration", "05-patient-registration.png", "အခြေခံ အချက်အလက် နှင့် G/P"),
                    ("Registration alerts", "05b-registration-alerts.png", "အသက် alert နှင့် warning ဥပမာ"),
                    ("Patient consent", "06-patient-consent.png", "လူနာသဘောတူချက်သည် အချက်အလက်အသုံးပြုမှုအတွက် လိုအပ်ပါသည်။"),
                ],
            },
        ],
    },
    {
        "title": "လူနာရှာဖွေခြင်းနှင့် Patient Care Hub",
        "sections": [
            {
                "heading": "လူနာရွေးခြင်း",
                "steps": [
                    "Home မှ Select Patient for Care ကိုနှိပ်ပါ။",
                    "Search box တွင် အမည်၊ ဖုန်း၊ patient ID ဖြင့် ရှာပါ။",
                    "Status chip များဖြင့် Registered, Antenatal, Labour, Postnatal, Baby လူနာများကို ခွဲရှာနိုင်ပါသည်။",
                    "မှန်ကန်သော လူနာ card ကိုနှိပ်ပါ။",
                    "Patient Care Hub တွင် ANC, Labour, Newborn, PNC, Immunization, Reports ကိုရွေးပါ။",
                ],
            },
            {
                "heading": "လူနာအချက်အလက် ပြင်ဆင်ခြင်း",
                "steps": [
                    "Patient list တွင် Edit ကိုနှိပ်ပါ။",
                    "ဖုန်း၊ လိပ်စာ၊ အခြေအနေစသည့် field များကိုသာ လိုအပ်သလို ပြင်ပါ။",
                    "Update/Save ကိုနှိပ်ပါ။",
                ],
            },
            {
                "heading": "Patient Care Hub",
                "paragraphs": [
                    "Patient Care Hub သည် လူနာတစ်ဦး၏ စောင့်ရှောက်မှုလုပ်ငန်းအားလုံးအတွက် အဓိကနေရာဖြစ်သည်။ Change Patient ဖြင့် လူနာပြောင်းနိုင်ပါသည်။",
                ],
                "screenshots": [
                    ("Patient list", "07-patient-list.png", "လူနာများကို search နှင့် filter ဖြင့် လျင်မြန်စွာရှာနိုင်ပါသည်။"),
                    ("Patient care hub", "08-patient-care-hub.png", "Care Hub မှ module အားလုံးကို ရွေးချယ်ပါ။"),
                ],
            },
        ],
    },
    {
        "title": "ANC — ကိုယ်ဝန်ဆောင် စောင့်ရှောက်မှု",
        "sections": [
            {
                "heading": "ANC Visit မှတ်တမ်းတင်ခြင်း",
                "steps": [
                    "Patient Care Hub မှ Antenatal Care ကိုနှိပ်ပါ။",
                    "New ANC Visit ကိုရွေးပါ။",
                    "Visit Date, LMP/EDD, gestational age (manual GA ထည့်နိုင်ပါသည်), history ကို စစ်ပါ။",
                    "Vital signs, BP, pulse, temperature, examination, danger signs, GBV စသည်တို့ကို ဖြည့်ပါ။",
                    "Urine protein/glucose အပေါ်လက်ရှိရလဒ်များတွင် သတိပေးချက် ပေါ်နိုင်ပါသည်။",
                    "အန္တရာယ်သတိပေးချက်ပေါ်လာပါက clinical guideline အတိုင်း ဆက်လုပ်ပါ။",
                    "Save Visit Data ကိုနှိပ်ပါ။",
                ],
            },
            {
                "heading": "MoH ၈-ကြိမ် ANC အချိန်ဇယား",
                "paragraphs": [
                    "App သည် MoH ၈-ကြိမ် ANC အချိန်ဇယားကို အကြံပြုပေးပါသည်။ Early ANC နှင့် manual gestational age ထည့်သွင်းမှု ပါဝင်ပါသည်။",
                ],
                "bullets": [
                    "ANC Report တွင် ၈-ကြိမ် visit register table ပုံစံဖြင့် ပရင့်ထုတ်နိုင်ပါသည်။",
                    "နောက်တစ်ကြိမ် ANC ရက်ကို form နှင့် report တွင် ကြည့်နိုင်ပါသည်။",
                ],
            },
            {
                "heading": "ANC Tests, Education, Report",
                "steps": [
                    "ANC Hub မှ Tests ကိုနှိပ်၍ lab test များကို ထည့်ပါ (HIV, Syphilis, Hep B/C, Ultrasound details စသည်)။",
                    "Education မှ Nutrition, Breastfeeding, Self-care, Pregnancy health သင်ခန်းစာများကို လူနာထံ ပြသနိုင်ပါသည်။",
                    "ANC Report မှ visit history, weight chart, lab results နှင့် next ANC date ကို ကြည့်ပါ။",
                    "Print ခလုတ်ဖြင့် A4 အစီရင်ခံစာ ထုတ်နိုင်ပါသည်။",
                ],
                "screenshots": [
                    ("ANC hub", "09-anc-hub.png", "ANC Hub မှ visit, report, tests, education ကိုရွေးပါ။"),
                    ("ANC form (top)", "10-anc-form-top.png", "Visit, LMP, EDD နှင့် early ANC badge"),
                    ("ANC form (vitals)", "10-anc-form-vitals.png", "Vitals, urine နှင့် danger signs"),
                    ("ANC report", "11-anc-report.png", "ANC Report တွင် visit အကျဉ်းချုပ်နှင့် timeline ကိုကြည့်ပါ။"),
                    ("ANC tests", "12-anc-tests.png", "Test result များကို ANC Tests တွင်သိမ်းပါ။"),
                ],
            },
        ],
    },
    {
        "title": "LCG — မွေးဖွားချိန် စောင့်ရှောက်မှု",
        "sections": [
            {
                "heading": "Labour Care စတင်ခြင်း",
                "steps": [
                    "Patient Care Hub မှ Labour Care ကိုနှိပ်ပါ။",
                    "ပထမဆုံးဝင်ပါက Labour setup တွင် labour onset, membrane, active first stage start time ကိုဖြည့်ပါ။",
                    "LCG entry တွင် အချိန်အလိုက် maternal condition, fetal condition, contractions, cervix စသည့်အချက်များကို ထည့်ပါ။",
                    "အချက်အလက်များကို မကြာခဏ Save လုပ်ပါ။",
                    "Summary View မှ partograph/chart ကိုကြည့်ပြီး လိုအပ်ပါက ပရင့်ထုတ်ပါ။",
                ],
                "callouts": [("danger", "Clinical alert", "အန္တရာယ်လက္ခဏာရှိပါက app ကိုမှတ်တမ်းတင်ရန်သာ မဟုတ်ဘဲ clinical protocol အတိုင်း ချက်ချင်း ဆောင်ရွက်ပါ။")],
            },
            {
                "heading": "Delivery Notes နှင့် Outcome",
                "steps": [
                    "မွေးဖွားပြီးပါက delivery notes တွင် မွေးဖွားချိန်၊ နည်းလမ်း၊ ကလေးအချက်အလက်များ ဖြည့်ပါ။",
                    "ကလေးမှတ်တမ်း အလိုအလျောက် ဖန်တီးနိုင်ပါသည်။",
                    "Other outcome သို့မဟုတ် transfer လိုပါက သက်ဆိုင်ရာ screen မှ ဖြည့်ပါ။",
                ],
                "screenshots": [
                    ("Labour setup", "13-labour-setup.png", "Labour setup သည် LCG မစတင်မီ အရေးကြီးပါသည်။"),
                    ("LCG entry", "14-lcg-entry.png", "LCG entry တွင် အချိန်လိုက် မှတ်တမ်းတင်ပါ။"),
                    ("LCG summary", "15-lcg-summary.png", "Summary View သည် chart ကြည့်ရန်နှင့် print ရန်ဖြစ်သည်။"),
                ],
            },
            {
                "heading": "Protocols နှင့် Emergencies",
                "paragraphs": [
                    "Labour Protocols နှင့် Labour Emergencies စာမျက်နှာများကို ကလিকယ် ကိုးကားချက်အဖြစ် အသုံးပြုနိုင်ပါသည်။ အမှန်တကယ် ဆောင်ရွက်ချက်များကို LCG entry နှင့် delivery notes တွင် မှတ်တမ်းတင်ပါ။",
                ],
            },
        ],
    },
    {
        "title": "မွေးကင်းစ စောင့်ရှောက်မှု",
        "sections": [
            {
                "heading": "Immediate Newborn Care",
                "steps": [
                    "Patient Care Hub မှ Newborn Care ကိုနှိပ်ပါ။",
                    "Immediate Newborn Care ကိုရွေးပါ။",
                    "Baby name, birth date/time, sex, weight, Apgar, immediate care procedures ကိုဖြည့်ပါ။",
                    "Vitamin K, eye care, cord care, breastfeeding support စသည်တို့ကို မှတ်တမ်းတင်ပါ။",
                    "Save ကိုနှိပ်ပါ။",
                ],
                "screenshots": [
                    ("Newborn hub", "17-newborn-hub.png", "Newborn Hub မှ immediate care, ongoing care, report ကိုရွေးပါ။"),
                    ("Newborn form (identity)", "18-newborn-identity.png", "Birth time, weight နှင့် locked fields"),
                    ("Newborn form (vitals/KMC)", "18-newborn-vitals-kmc.png", "Vitals, feeding နှင့် KMC section"),
                ],
            },
            {
                "heading": "Newborn Care Visits (အကြိမ် ၁–၄)",
                "steps": [
                    "Newborn Care Page တွင် visit 1, 2, 3, 4 ကို မွေးပြီး ၀၊ ၃၊ ၁၄၊ ၄၂ ရက် အချိန်ဇယားအတိုင်း မှတ်တမ်းတင်ပါ။",
                    "Vitals, Respiration, Feeding — ကလေးအားလုံးအတွက် ဖြည့်ပါ။",
                    "Assessment — မျက်လုံး၊ ချက်ကြိုး၊ immunization modal ဖြင့် ကာကွယ်ဆေး ထည့်နိုင်ပါသည်။",
                    "KMC — ကိုယ်အလေးချိန် ၂၀၀၀–၂၅၀၀ g ဖြစ်ပါက potential KMC အဖြစ် ဆုံးဖြတ်ချက် ဖြည့်နိုင်ပါသည်။",
                    "Outcome နှင့် Clinical notes ကို visit အဆုံးတွင် ဖြည့်ပြီး Save Newborn Care Data ကိုနှိပ်ပါ။",
                ],
            },
            {
                "heading": "KMC Tracking",
                "steps": [
                    "KMC ဆုံးဖြတ်ချက် Yes ဖြစ်ပါက KMC hours, discharge date ဖြည့်ပါ။",
                    "Home သို့မဟုတ် KMC Tracking screen မှ follow-up ကြည့်နိုင်ပါသည်။",
                    "Newborn Report တွင် KMC အချက်အလက် ပါဝင်ပါသည် (KMC baby ဖြစ်မှသာ)။",
                ],
            },
            {
                "heading": "Newborn Report",
                "paragraphs": [
                    "Newborn Report တွင် visit tabs ဖြင့် ကြည့်နိုင်ပြီး Print လုပ်ပါက visit register table ပုံစံဖြင့် အကြိမ် ၁–၄ အားလုံး ပါဝင်ပါသည်။",
                ],
                "screenshots": [("Newborn report", "19-newborn-report.png", "Newborn Report တွင် ကလေးမှတ်တမ်းကို ကြည့်ပြီး print ထုတ်ပါ။")],
            },
        ],
    },
    {
        "title": "PNC — မွေးပြီးမိခင် ပြုစုမှု",
        "sections": [
            {
                "heading": "PNC Visit မှတ်တမ်းတင်ခြင်း",
                "steps": [
                    "Patient Care Hub မှ PNC သို့မဟုတ် Postnatal Care ကိုနှိပ်ပါ။",
                    "New PNC Visit ကိုနှိပ်ပါ။",
                    "Delivery date/time သည် newborn visit 1 သို့မဟုတ် delivery notes နှင့် မျှဝေထားနိုင်ပါသည် — တစ်နေရာတည်းမှ ထည့်ပြီးပါက နှစ်ခုလုံး lock ဖြစ်နိုင်ပါသည်။",
                    "Vital signs, physical examination, danger signs, treatment, notes ကိုဖြည့်ပါ။",
                    "လွှဲပြောင်းရန်လိုပါက Refer Patient ကိုနှိပ်ပါ။",
                    "Save Postpartum Visit ကိုနှိပ်ပါ။",
                ],
                "bullets": [
                    "PNC ၄-ကြိမ် အချိန်ဇယားသည် newborn ၄-ကြိမ် အချိန်ဇယားနှင့် တူညီပါသည် (+၃, +၁၄, +၄၂ ရက်)။",
                ],
                "screenshots": [
                    ("PNC hub", "20-pnc-hub.png", "PNC Hub မှ visit အသစ်နှင့် PNC report ကိုရွေးပါ။"),
                    ("PNC form (top)", "21-pnc-form-top.png", "Visit, delivery anchor နှင့် vitals"),
                    ("PNC form (exam)", "21-pnc-form-exam.png", "Examination, danger signs, outcome"),
                    ("PNC report", "22-pnc-report.png", "PNC Report တွင် visit history ကို ကြည့်ပါ။"),
                ],
            },
        ],
    },
    {
        "title": "ကာကွယ်ဆေး",
        "sections": [
            {
                "heading": "Immunization Hub",
                "steps": [
                    "Baby patient ၏ Patient Care Hub မှ Immunization ကိုနှိပ်ပါ။",
                    "Myanmar Routine Immunization Schedule တွင် Due soon / Given / Pending အခြေအနေကို ကြည့်ပါ။",
                    "Record Vaccine ဖြင့် ကာကွယ်ဆေးမှတ်တမ်း အသစ် ထည့်ပါ။",
                    "Newborn care page မှ immunization modal ဖြင့် Hep B, BCG စသည်တို့ကိုလည်း တိုက်ရိုက်သိမ်းနိုင်ပါသည်။",
                    "နောက်တစ်ကြိမ်လာရမည့်ရက်ကို မိသားစုအား ရှင်းပြပါ။",
                ],
                "screenshots": [
                    ("Vaccine home", "23-vaccine-home.png", "ကာကွယ်ဆေး Hub မှ schedule နှင့် saved records ကိုကြည့်ပါ။"),
                    ("Vaccine record", "24-vaccine-record.png", "ကာကွယ်ဆေးထိုးမှတ်တမ်းကို သိမ်းပါ။"),
                ],
            },
        ],
    },
    {
        "title": "လွှဲပြောင်းခြင်းနှင့် Referral",
        "sections": [
            {
                "heading": "Internal Transfer",
                "steps": [
                    "Care form မှ Refer Patient သို့မဟုတ် Transfer screen ကိုဖွင့်ပါ။",
                    "လွှဲပြောင်းရသည့် facility ကို ရွေးပါ။",
                    "အကြောင်းရင်း နှင့် clinical note ဖြည့်ပါ။",
                    "Submit/Save လုပ်ပါ။",
                ],
            },
            {
                "heading": "External Transfer",
                "steps": [
                    "Transfer Patient screen တွင် External ကိုရွေးပါ။",
                    "ပြင်ပ ဆေးရုံ/ဌာန အမည် ဖြည့်ပါ။",
                    "လူနာအခြေအနေ referred_external အဖြစ် ပြင်ဆင်နိုင်ပါသည်။",
                ],
                "screenshots": [("Transfer screen", "16-transfer.png", "လွှဲပြောင်းရန်လိုပါက Transfer form ကိုသုံးပါ။")],
            },
            {
                "heading": "Patient Transfers (Home)",
                "paragraphs": [
                    "Home မှ Patient Transfers card ဖြင့် ဝင်လာသော transfer request များကို လက်ခံနိုင်ပါသည်။",
                ],
            },
        ],
    },
    {
        "title": "အစီရင်ခံစာများနှင့် ပရင့်ထုတ်ခြင်း",
        "sections": [
            {
                "heading": "Reports ကြည့်ခြင်း",
                "steps": [
                    "Patient Care Hub မှ individual report သို့မဟုတ် Overall Patient Report ကိုရွေးပါ။",
                    "Overall Patient Report တွင် Overview, ANC, Labour, Newborn, PNC tabs များရှိသည်။",
                    "ANC Report, Newborn Report, PNC Report, Lab Report တို့ကိုလည်း သီးခြားကြည့်နိုင်ပါသည်။",
                    "မြန်မာ/အင်္ဂလိပ် ပြောင်းလိုပါက ENG/MM ကိုနှိပ်ပါ။",
                    "ပရင့်ထုတ်ရန် Print ခလုတ်ကိုနှိပ်ပါ။ Phone တွင် PDF အဖြစ် Save လုပ်နိုင်ပါသည်။",
                ],
                "screenshots": [
                    ("Overall patient report", "25-overall-report.png", "Overall Report သည် လူနာတစ်ဦး၏ care history အားလုံးကိုစုစည်းပေးသည်။"),
                    ("Print report", "26-print-report.png", "Print/PDF ကို official record အတွက် အသုံးပြုနိုင်ပါသည်။"),
                ],
            },
        ],
    },
    {
        "title": "Settings, CME နှင့် အခြား Module များ",
        "sections": [
            {
                "heading": "Settings",
                "steps": [
                    "Home မှ Settings (gear icon) ကိုနှိပ်ပါ။",
                    "ဘာသာစကား၊ password ပြောင်းခြင်း စသည်တို့ကို လုပ်ဆောင်နိုင်ပါသည်။",
                ],
            },
            {
                "heading": "CME",
                "steps": [
                    "Home မှ CME - Continuous Medical Learning ကိုနှိပ်ပါ။",
                    "Mandatory နှင့် Optional module များကို ဖတ်ပြီး ပြီးမြောက်အောင် လုပ်ပါ။",
                ],
            },
            {
                "heading": "Dashboard နှင့် High Risk",
                "paragraphs": [
                    "Analytics Dashboard — facility အဆင့် စာရင်းဇယားများကြည့်ရန် (internet လိုအပ်နိုင်ပါသည်)။",
                    "High Risk Tracking — follow-up လိုသော အန္တရာယ်မြင့် လူနာများစာရင်း။",
                    "Township Report — TMO အတွက် မြို့နယ်အစီရင်ခံစာ။",
                ],
            },
            {
                "heading": "Feedback နှင့် Logout",
                "steps": [
                    "Pilot အတွင်း Feedback form ဖြင့် အကြံပြုချက်ပေးနိုင်ပါသည်။",
                    "အလုပ်ပြီးပါက Logout နှိပ်ပြီး shared device တွင် လုံခြုံရေးထိန်းပါ။",
                ],
            },
        ],
    },
    {
        "title": "Offline နှင့် Sync",
        "sections": [
            {
                "heading": "Internet မရှိသောအချိန်",
                "steps": [
                    "Login ဝင်စဉ် Keep me logged in ကို အမှန်ခြစ်ထားပါက offline တွင် အချို့လုပ်ငန်းများ ဆက်လုပ်နိုင်ပါသည်။",
                    "Offline banner ပေါ်လာပါက Internet မရှိကြောင်း သိနိုင်ပါသည်။",
                    "လူနာမှတ်တမ်းများကို local device ထဲတွင် ယာယီသိမ်းထားပါသည်။",
                    "Internet ပြန်ရပါက Home တွင် Sync ကိုနှိပ်ပါ။",
                    "Sync ပြီးကြောင်း message ပေါ်လာသည်အထိ စောင့်ပါ။",
                ],
                "callouts": [("warning", "အရေးကြီး", "Sync မလုပ်မီ app ကို မဖျက်ပါနှင့်။ Phone browser data မရှင်းပါနှင့်။")],
                "screenshots": [("Offline and sync", "27-offline-sync.png", "Home တွင် Offline/Sync အခြေအနေကို စစ်နိုင်ပါသည်။")],
            },
        ],
    },
    {
        "title": "ပြဿနာဖြေရှင်းနည်း",
        "sections": [
            {
                "heading": "အဖြစ်များသော ပြဿနာများ",
                "table": {
                    "headers": ["ပြဿနာ", "လုပ်ဆောင်ရန်"],
                    "rows": [
                        ["Login မဝင်နိုင်ပါ", "Internet ရှိ/မရှိ စစ်ပါ။ Email/Password မှန်မမှန် စစ်ပါ။ အကောင့် approval ရပြီးမရပြီးကို supervisor ကိုမေးပါ။"],
                        ["လူနာမတွေ့ပါ", "အမည် spelling, ဖုန်းနံပါတ်, patient ID ဖြင့် ပြန်ရှာပါ။ Status filter ကို All သို့ပြောင်းပါ။ Baby/Mommy tab စစ်ပါ။"],
                        ["Save မလုပ်နိုင်ပါ", "Required field များပြည့်စုံမစုံ စစ်ပါ။ Red warning message ကိုဖတ်ပါ။ Internet မရှိပါက local save ဖြစ်နိုင်သည်။"],
                        ["App ပြောင်းလဲမှုအသစ် မပေါ်ပါ", "Hard refresh လုပ်ပါ။ PWA ဖြစ်ပါက app ကိုပိတ်ပြီးပြန်ဖွင့်ပါ။"],
                        ["Sync မအောင်မြင်ပါ", "Internet တည်ငြိမ်အောင်စောင့်ပါ။ Home သို့ပြန်သွားပြီး Sync ကို ထပ်နှိပ်ပါ။ App data ကို မဖျက်ပါနှင့်။"],
                        ["ကာကွယ်ဆေး schedule မပြင်ပါ", "vaccineId မှန်ကန်စွာ သိမ်းထားကြောင်း စစ်ပါ။ Page refresh လုပ်ပါ။"],
                    ],
                },
                "callouts": [("success", "နေ့စဉ်အဆုံးတွင်", "လူနာမှတ်တမ်းများ Save ဖြစ်မဖြစ် စစ်ပါ၊ Internet ရှိပါက Sync လုပ်ပါ၊ Logout မလုပ်မီ pending records မရှိကြောင်း စစ်ပါ။")],
            },
        ],
    },
]


def build_document():
    global _chapter_heading_count
    _chapter_heading_count = 0
    manual = apply_detailed_sections(MANUAL)
    doc = Document()

    add_cover_page(doc)
    add_title_page(doc)
    add_front_matter_page(doc)

    # TOC (still in front-matter section — no page numbers)
    add_heading(doc, "မာတိကာ", level=1)
    for i, ch in enumerate(manual, 1):
        add_para(doc, f"{i}. {ch['title']}", size=11)
    doc.paragraphs[-1].runs[0].add_break(WD_BREAK.PAGE)

    # Body section: A5 landscape, binding gutter, header/footer, page numbers from 1
    body_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    clear_header_footer(body_section)
    configure_section(body_section, landscape=True, header_footer=True)
    start_page_numbering_at_one(body_section)

    # Chapters
    for i, chapter in enumerate(manual, 1):
        add_para(doc, f"အခန်း {i}", size=10, color=RGBColor(0x64, 0x74, 0x8B))
        add_heading(doc, chapter["title"], level=1, page_break_before=True)
        for sec in chapter.get("sections", []):
            if sec.get("heading"):
                add_heading(doc, sec["heading"], level=2)
            for para in sec.get("paragraphs", []):
                add_para(doc, para)
            for item in sec.get("bullets", []):
                add_bullets(doc, [item])
            if sec.get("steps"):
                add_steps(doc, sec["steps"])
            if sec.get("field_table"):
                ft = sec["field_table"]
                add_field_table(doc, ft["headers"], ft["rows"])
            if sec.get("table"):
                add_table(doc, sec["table"]["headers"], sec["table"]["rows"])
            for kind, label, text in sec.get("callouts", []):
                add_callout(doc, label, text, kind)
            if sec.get("screenshot_rows"):
                for row in sec["screenshot_rows"]:
                    add_screenshot_row(doc, row, cols=sec.get("screenshot_cols", 2))
            elif sec.get("screenshots"):
                add_screenshot_row(doc, sec["screenshots"], cols=sec.get("screenshot_cols", 2))

    add_para(doc, "— အဆုံးသတ် —", size=11, bold=True, space_after=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "m-MNCH Care | ကျန်းမာရေးဝန်ကြီးဌာန လုပ်ငန်းစဉ်များ | Jhpiego Myanmar",
        size=9,
        color=RGBColor(0x64, 0x74, 0x8B),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )

    return doc


def main():
    doc = build_document()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    doc.save(OUT_PATH_ALT)
    print(f"Wrote {OUT_PATH}")
    print(f"Wrote {OUT_PATH_ALT}")


if __name__ == "__main__":
    main()
